# requirements: streamlit, python-docx, PyPDF2
import streamlit as st
from docx import Document
from io import BytesIO
import os
import PyPDF2

def extract_text_from_docx(file_bytes):
    """Extract text from a DOCX file"""
    try:
        doc = Document(BytesIO(file_bytes))
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)
    except Exception as e:
        return f"Error reading DOCX file: {str(e)}"

def extract_text_from_pdf(file_bytes):
    """Extract text from a PDF file"""
    try:
        pdf_reader = PyPDF2.PdfReader(BytesIO(file_bytes))
        full_text = []
        for page in pdf_reader.pages:
            full_text.append(page.extract_text())
        return '\n'.join(full_text)
    except Exception as e:
        return f"Error reading PDF file: {str(e)}"

def get_document_stats(doc):
    """Get basic statistics from a Document object"""
    try:
        paragraph_count = len(doc.paragraphs)
        word_count = sum(len(p.text.split()) for p in doc.paragraphs)
        char_count = sum(len(p.text) for p in doc.paragraphs)
        return paragraph_count, word_count, char_count
    except:
        return 0, 0, 0

def get_pdf_stats(text):
    """Get basic statistics from PDF text"""
    try:
        lines = text.split('\n')
        words = text.split()
        return len(lines), len(words), len(text)
    except:
        return 0, 0, 0

def create_document_from_text(text):
    """Create a Word document from text content"""
    doc = Document()
    paragraphs = text.split('\n')
    for para in paragraphs:
        doc.add_paragraph(para)
    return doc

def document_to_bytes(doc):
    """Convert Document object to bytes"""
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Page configuration
st.set_page_config(
    page_title="Document Reader",
    page_icon="📄",
    layout="wide"
)

# Title and description
st.title("📄 Multi-Document Reader")
st.markdown("""
Upload, read, and create Microsoft Word (.docx) and PDF documents:
- Extract and view document contents
- See basic document statistics
- Create and edit new documents
- Download in various formats
""")

# File uploader
uploaded_files = st.file_uploader(
    "Choose documents (Word or PDF)",
    type=["docx", "pdf"],
    accept_multiple_files=True,
    key="uploader"
)

# Create tabs
tabs = st.tabs(["📁 Uploaded Documents", "➕ Create New Document"])

# Tab 1: Uploaded Documents
with tabs[0]:
    if uploaded_files:
        st.subheader("Document Contents")
        
        # Create tabs for each document
        doc_tabs = st.tabs([f"📄 {f.name}" for f in uploaded_files])
        
        for i, (file, doc_tab) in enumerate(zip(uploaded_files, doc_tabs)):
            with doc_tab:
                # Extract text based on file type
                file_bytes = file.getvalue()
                file_extension = os.path.splitext(file.name)[1].lower()
                
                if file_extension == ".docx":
                    text_content = extract_text_from_docx(file_bytes)
                elif file_extension == ".pdf":
                    text_content = extract_text_from_pdf(file_bytes)
                else:
                    text_content = "Unsupported file type"
                
                # Display stats
                if file_extension == ".docx":
                    try:
                        doc = Document(BytesIO(file_bytes))
                        p_count, w_count, c_count = get_document_stats(doc)
                        
                        col1, col2, col3 = st.columns(3)
                        col1.metric("Paragraphs", p_count)
                        col2.metric("Words", w_count)
                        col3.metric("Characters", c_count)
                    except:
                        st.warning("Could not calculate document statistics")
                elif file_extension == ".pdf":
                    p_count, w_count, c_count = get_pdf_stats(text_content)
                    
                    col1, col2, col3 = st.columns(3)
                    col1.metric("Lines", p_count)
                    col2.metric("Words", w_count)
                    col3.metric("Characters", c_count)
                
                # Display content
                edited_content = st.text_area(
                    "Document Text (Editable)",
                    value=text_content,
                    height=300,
                    key=f"edit_area_{i}"
                )
                
                # Download buttons
                st.markdown("---")
                st.subheader("Save Options")
                
                if file_extension == ".docx":
                    # Create updated DOCX from edited content
                    updated_doc = create_document_from_text(edited_content)
                    updated_bytes = document_to_bytes(updated_doc)
                    
                    col1, col2, col3 = st.columns(3)
                    col1.download_button(
                        label="📄 Save as DOCX",
                        data=updated_bytes,
                        file_name=f"{os.path.splitext(file.name)[0]}_edited.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        help="Downloads to your browser's default download folder"
                    )
                    
                    col2.download_button(
                        label="💾 Save as Text",
                        data=edited_content,
                        file_name=f"{os.path.splitext(file.name)[0]}_edited.txt",
                        mime="text/plain",
                        help="Downloads to your browser's default download folder"
                    )
                    
                    col3.download_button(
                        label="🖨️ Save as PDF",
                        data=edited_content,
                        file_name=f"{os.path.splitext(file.name)[0]}_edited.txt",
                        mime="text/plain",
                        help="Save as text then print to PDF"
                    )
                
                else:  # PDF
                    # For PDFs, we can only save as text or convert to DOCX
                    new_doc = create_document_from_text(edited_content)
                    new_doc_bytes = document_to_bytes(new_doc)
                    
                    col1, col2 = st.columns(2)
                    col1.download_button(
                        label="📄 Save as DOCX",
                        data=new_doc_bytes,
                        file_name=f"{os.path.splitext(file.name)[0]}_converted.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        help="Converts PDF content to Word document"
                    )
                    
                    col2.download_button(
                        label="💾 Save as Text",
                        data=edited_content,
                        file_name=f"{os.path.splitext(file.name)[0]}_converted.txt",
                        mime="text/plain",
                        help="Extracted text from PDF"
                    )
                
                st.info("📁 Files save to your browser's default download location")
    else:
        st.info("Please upload .docx or .pdf files to get started")

# Tab 2: Create New Document
with tabs[1]:
    st.subheader("Create and Edit New Document")
    
    # Initialize session state for document content
    if "doc_content" not in st.session_state:
        st.session_state.doc_content = "# My Document\n\nStart typing here...\n\n- Item 1\n- Item 2\n\n**Bold text** and *italic text*"
    
    # Text area for document content
    st.session_state.doc_content = st.text_area(
        "Document Content (Markdown-style supported)",
        value=st.session_state.doc_content,
        height=400,
        key="doc_editor"
    )
    
    # Instructions
    with st.expander("💡 Formatting Tips"):
        st.markdown("""
        - **Headings**: Start lines with `#` (H1), `##` (H2), etc.
        - **Bold**: Wrap text with `**bold**`
        - **Italic**: Wrap text with `*italic*`
        - **Lists**: Use `- ` for bullet points
        - **New Paragraphs**: Separate with blank lines
        """)
    
    # Preview section
    st.subheader("Preview")
    st.markdown(st.session_state.doc_content)
    
    # Create document from content
    doc = create_document_from_text(st.session_state.doc_content)
    doc_bytes = document_to_bytes(doc)
    
    # Download buttons
    st.markdown("---")
    st.subheader("Save Your Document")
    
    col1, col2, col3 = st.columns(3)
    col1.download_button(
        label="📄 Save as DOCX",
        data=doc_bytes,
        file_name="my_document.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        help="Standard Word document format"
    )
    
    col2.download_button(
        label="💾 Save as Text",
        data=st.session_state.doc_content,
        file_name="my_document.txt",
        mime="text/plain",
        help="Plain text format"
    )
    
    col3.download_button(
        label="🖨️ Save as PDF",
        data=st.session_state.doc_content,
        file_name="my_document.txt",
        mime="text/plain",
        help="Save as text then print to PDF"
    )
    
    st.info("📁 Files save to your browser's default download location")

# Additional information
with st.expander("ℹ️ How to use this tool"):
    st.markdown("""
    1. **Supported Formats**:
       - Microsoft Word (.docx files)
       - PDF Documents (.pdf files)
    2. **Upload Documents**: 
       - Use the file uploader to select documents
       - View and edit document content in text areas
       - Save modifications using the download buttons
    3. **Create New Documents**:
       - Switch to the "Create New Document" tab
       - Edit content directly in the text area
       - See real-time preview of your document
       - Save when finished using the save buttons
    4. **File Saving**:
       - Files download directly to your browser's default download location
       - Check your Downloads folder or browser download bar
       - Some browsers may ask before downloading multiple files
    """)
    
    st.markdown("""
    **Note on PDF Handling**: 
    - PDF text extraction may not preserve original formatting
    - Complex layouts, images, and tables may not be fully extracted
    - Converted PDFs become editable text documents
    """)
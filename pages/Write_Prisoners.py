# requirements: streamlit, python-docx, PyPDF2, reportlab
import streamlit as st
from docx import Document
from io import BytesIO
import os
import PyPDF2
import json
import base64
import time
import pandas as pd
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch

# Add import for our search widget
from utils.search_widget import render_search_widget


def extract_text_from_docx(file_bytes):
    """Extract text from a DOCX file"""
    try:
        doc = Document(BytesIO(file_bytes))
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)
    except Exception as e:
        return f"Error reading DOCX file: {str(e)}"

def extract_text_from_pdf(file_bytes):
    """Extract text from a PDF file"""
    try:
        pdf_reader = PyPDF2.PdfReader(BytesIO(file_bytes))
        full_text = []
        for page in pdf_reader.pages:
            full_text.append(page.extract_text())
        return '\n'.join(full_text)
    except Exception as e:
        return f"Error reading PDF file: {str(e)}"

def get_document_stats(doc):
    """Get basic statistics from a Document object"""
    try:
        paragraph_count = len(doc.paragraphs)
        word_count = sum(len(p.text.split()) for p in doc.paragraphs)
        char_count = sum(len(p.text) for p in doc.paragraphs)
        return paragraph_count, word_count, char_count
    except:
        return 0, 0, 0

def get_pdf_stats(text):
    """Get basic statistics from PDF text"""
    try:
        lines = text.split('\n')
        words = text.split()
        return len(lines), len(words), len(text)
    except:
        return 0, 0, 0

def create_document_from_text(text):
    """Create a Word document from text content"""
    doc = Document()
    paragraphs = text.split('\n')
    for para in paragraphs:
        doc.add_paragraph(para)
    return doc

def document_to_bytes(doc):
    """Convert Document object to bytes"""
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def create_pdf_from_text(text):
    """Create a PDF document from text content"""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    
    # Create a custom style for better formatting
    custom_style = ParagraphStyle(
        'Custom',
        parent=styles['Normal'],
        fontSize=12,
        spaceAfter=12,
    )
    
    # Split text into paragraphs
    paragraphs = text.split('\n')
    story = []
    
    for para in paragraphs:
        if para.strip():  # Only add non-empty paragraphs
            p = Paragraph(para, custom_style)
            story.append(p)
        else:
            # Add spacing for empty lines
            story.append(Spacer(1, 0.2*inch))
    
    doc.build(story)
    buffer.seek(0)
    return buffer

def save_session(uploaded_files):
    """Save current session state to a file"""
    # Get information about all uploaded files
    uploaded_files_info = []
    if uploaded_files:
        for file in uploaded_files:
            uploaded_files_info.append({
                "name": file.name,
                "size": len(file.getvalue()) if hasattr(file, 'getvalue') else 0,
                "type": os.path.splitext(file.name)[1].lower() if '.' in file.name else 'unknown'
            })
    
    # Collect all document content from session state
    doc_contents = {}
    for key in st.session_state:
        if key.startswith("doc_content_"):
            doc_contents[key] = st.session_state[key]
    
    session_data = {
        "timestamp": time.time(),
        "selected_doc": st.session_state.get("selected_doc", None),
        "doc_content": st.session_state.get("doc_content", ""),
        "doc_contents": doc_contents,  # Include all document contents
        "uploaded_files": uploaded_files_info
    }
    
    # Create sessions directory if it doesn't exist
    if not os.path.exists("sessions"):
        os.makedirs("sessions")
    
    # Save session data to file
    session_file = f"sessions/session_{int(time.time())}.json"
    with open(session_file, "w") as f:
        json.dump(session_data, f)
    
    return session_file

def load_session(session_file):
    """Load session state from a file"""
    try:
        with open(session_file, "r") as f:
            session_data = json.load(f)
        
        # Restore session state
        st.session_state.selected_doc = session_data.get("selected_doc", None)
        st.session_state.doc_content = session_data.get("doc_content", "")
        
        # Restore document contents
        doc_contents = session_data.get("doc_contents", {})
        for key, content in doc_contents.items():
            st.session_state[key] = content
        
        return True
    except Exception as e:
        st.error(f"Error loading session: {str(e)}")
        return False

def get_session_files():
    """Get list of available session files"""
    if not os.path.exists("sessions"):
        return []
    
    session_files = [f for f in os.listdir("sessions") if f.endswith(".json")]
    return sorted(session_files, reverse=True)  # Most recent first

def render_write_prisoners():
    """Render the Word/PDF editor with card-based layout"""
    # Page configuration
    st.set_page_config(
        page_title="Document Reader - Cards",
        page_icon="📄",
        layout="wide"
    )
    
    # Initialize last save time in session state
    if "last_save_time" not in st.session_state:
        st.session_state.last_save_time = time.time()
    
    # Auto-restore the most recent session if no document is currently selected
    # This helps prevent data loss when switching tabs or reloading the page
    if "selected_doc" not in st.session_state or st.session_state.selected_doc is None:
        session_files = get_session_files()
        if session_files:
            # Try to auto-restore the most recent session
            most_recent_session = session_files[0]  # First one is most recent due to sorting
            try:
                with open(f"sessions/{most_recent_session}", "r") as f:
                    session_data = json.load(f)
                
                # Only auto-restore if the session has actual content
                if session_data.get("doc_content") and session_data.get("doc_content") != "# My Document\n\nStart typing here...\n\n- Item 1\n- Item 2\n\n**Bold text** and *italic text*":
                    # Auto-restore the session
                    load_session(f"sessions/{most_recent_session}")
                    # Don't show the manual restore UI if we've auto-restored
                    session_files = []  # Clear the list so we don't show the restore UI
            except Exception as e:
                # Silently fail to avoid disrupting user experience
                pass
    
    # Check for existing sessions and offer to restore
    # Always check for session files, not just when we think it's a fresh session
    session_files = get_session_files()
    if session_files:
        # Collect session data for all available sessions
        session_options = []
        session_details = {}
        
        for session_file in session_files:
            try:
                with open(f"sessions/{session_file}", "r") as f:
                    session_data = json.load(f)
                
                # Only include sessions with actual data
                if session_data.get("selected_doc") or \
                   (session_data.get("doc_content") and session_data.get("doc_content") != "# My Document\n\nStart typing here...\n\n- Item 1\n- Item 2\n\n**Bold text** and *italic text*"):
                    # Create a display name for the session
                    timestamp = session_data.get("timestamp", 0)
                    doc_name = session_data.get("selected_doc", "New Document")
                    display_name = f"{time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(timestamp))} - {doc_name}"
                    
                    session_options.append(display_name)
                    session_details[display_name] = {
                        "file": session_file,
                        "data": session_data
                    }
            except Exception as e:
                # Silently fail to avoid disrupting user experience
                pass
        
        # Show session restoration UI if we have sessions with data
        if session_options:
            with st.sidebar:
                st.markdown("---")
                st.subheader("💾 Previous Sessions")
                st.caption("Select a previous session to restore:")
                
                # Create a selectbox for session selection
                selected_session = st.selectbox(
                    "Choose session to restore:",
                    session_options,
                    key="session_selector"
                )
                
                if selected_session and selected_session in session_details:
                    session_info = session_details[selected_session]
                    st.caption(f"Document: {session_info['data'].get('selected_doc', 'New Document')}")
                    st.caption(f"Saved: {time.ctime(session_info['data'].get('timestamp', 0))}")
                    
                    # Show list of files from the previous session
                    uploaded_files_info = session_info['data'].get('uploaded_files', [])
                    if uploaded_files_info:
                        st.markdown("**Files in this session:**")
                        for file_info in uploaded_files_info:
                            st.markdown(f"- {file_info['name']} ({file_info['type']})")
                    elif session_info['data'].get('selected_doc'):
                        # For older session files that don't have uploaded_files info
                        st.markdown("**Selected document in this session:**")
                        st.markdown(f"- {session_info['data'].get('selected_doc')}")
                    
                    # Show information about the session
                    st.info("💡 To fully restore this session, please re-upload the files listed above.")
                    
                    if st.button("Restore Selected Session"):
                        load_session(f"sessions/{session_info['file']}")
                        st.rerun()
    
    # Add a manual save session button in the sidebar (always visible)
    with st.sidebar:
        st.markdown("---")
        st.markdown("Ctrl-Enter New Document before Saving")
        if st.button("💾 Save Current Session"):
            try:
                # Save session with empty list to capture all session data
                save_session([])
                st.success("Session saved successfully!")
            except Exception as e:
                st.error(f"Error saving session: {str(e)}")
    
    # Auto-save every 5 seconds (increased frequency for better data protection)
    current_time = time.time()
    if current_time - st.session_state.last_save_time > 5:
        try:
            # Save session with uploaded_files if available, otherwise save with empty list
            # This ensures we always save session data even if uploaded_files is not in scope
            if 'uploaded_files' in locals():
                save_session(uploaded_files)
            else:
                # Create a mock uploaded_files list from session state to ensure all content is saved
                mock_uploaded_files = []
                # Add any files that have content in session state
                for key in st.session_state:
                    if key.startswith("doc_content_"):
                        # Extract filename from key
                        filename = key.replace("doc_content_", "")
                        # Create a mock file object with just the name
                        class MockFile:
                            def __init__(self, name):
                                self.name = name
                        mock_uploaded_files.append(MockFile(filename))
                save_session(mock_uploaded_files)
            st.session_state.last_save_time = current_time
        except Exception as e:
            # Log the error for debugging (this will appear in the console)
            print(f"Auto-save error: {str(e)}")
            pass  # Silently fail to avoid disrupting user experience

    # Title and description
    st.title("📄 Writing to Prisoners (Card UI)")

    # Add search functionality for main prisoner data
    st.markdown("---")
    st.markdown("### 🔍 Prisoner Search, find CPID")


    
    # Check if prisoner data is available
    if 'df' in st.session_state and not st.session_state.df.empty:
        # Allow user to select columns to display
        all_columns = st.session_state.df.columns.tolist()
        default_columns = ['fName', 'lName', 'Sponsor', 'CPID','Step (received only)', 'letter exchange (received only)']
        
        # Filter default_columns to only include columns that exist in the dataframe
        existing_default_columns = [col for col in default_columns if col in all_columns]
        
        # Create a multiselect for column selection
        selected_columns = st.multiselect(
            "Select columns to display in search results:",
            options=all_columns,
            default=existing_default_columns
        )
        
        # Use our search widget to search the main data
        search_results = render_search_widget(
            df=st.session_state.df,
            search_column='lName',
            display_columns=selected_columns if selected_columns else all_columns,
            search_label="Search by Last Name",
            button_label="Search Records"
        )
        
        if search_results is not None and not search_results.empty:
            #Confidential notification
            st.markdown("""
            <div style='text-align: center;'>
                <span style='color: red; font-size: 24px; font-weight: bold;'>CONFIDENTIAL PERSONAL INFO</span>
            </div>
            """, unsafe_allow_html=True)

            # We could add additional actions here for the search results
    else:
        st.info("Prisoner database not loaded. Please upload data on the home page.")

    st.markdown("---")

    # Import directory selection widget
    from utils.directory_selection_widget import directory_selection_widget

    # 📁 Directory Selection
    directory_selection_widget()

    # File uploader
    uploaded_files = st.file_uploader(
        "Upload documents (Word or PDF)",
        type=["docx", "pdf"],
        accept_multiple_files=True,
        key="uploader"
    )


    # Add search functionality for document names
    if uploaded_files:
        st.markdown("---")
        st.markdown("### 🔍 Search Uploaded Documents")
        
        # Create a dataframe with document information for searching
        doc_data = []
        for file in uploaded_files:
            doc_data.append({
                'name': file.name,
                'type': os.path.splitext(file.name)[1].lower(),
                'size': f"{len(file.getvalue()) / 1024:.1f} KB"
            })
        
        doc_df = pd.DataFrame(doc_data)
        
        # Use our search widget to search document names
        doc_search_results = render_search_widget(
            df=doc_df,
            search_column='name',
            display_columns=['name', 'type', 'size'],
            search_label="Search Documents by Name",
            button_label="Find Documents"
        )
        
        if doc_search_results is not None and not doc_search_results.empty:
            st.markdown("#### Document Search Results:")
            # We could add additional actions here for the search results

    # Initialize session state for selected document
    if "selected_doc" not in st.session_state:
        st.session_state.selected_doc = None

    # Create tabs for main content area
    tabs = st.tabs(["📁 Uploaded Documents", "➕ Create New Document"])

    # Tab 1: Uploaded Documents (Card-based layout)
    with tabs[0]:
        if uploaded_files:
            st.subheader("Document Cards")
            st.markdown("Click on any document card to view and edit its content.")
            
            # Create a grid of cards
            cols = st.columns(3)  # 3 columns for card grid
            
            # Display each document as a card
            for i, file in enumerate(uploaded_files):
                with cols[i % 3]:
                    # Card container
                    with st.container(border=True):
                        # Document icon and name
                        file_extension = os.path.splitext(file.name)[1].lower()
                        icon = "📄" if file_extension == ".docx" else "📑" if file_extension == ".pdf" else "📁"
                        st.markdown(f"### {icon} {file.name}")
                        
                        # File size
                        file_size_kb = len(file.getvalue()) / 1024
                        st.caption(f"Size: {file_size_kb:.1f} KB")
                        
                        # Preview of document content (first 100 characters)
                        file_bytes = file.getvalue()
                        if file_extension == ".docx":
                            preview_text = extract_text_from_docx(file_bytes)[:100] + "..."
                        elif file_extension == ".pdf":
                            preview_text = extract_text_from_pdf(file_bytes)[:100] + "..."
                        else:
                            preview_text = "Unsupported file type"
                        st.caption(f"Preview: {preview_text}")
                        
                        # Action button to select this document
                        if st.button("View/Edit", key=f"select_{file.name}"):
                            st.session_state.selected_doc = file.name
                            st.rerun()
            
            # If a document is selected, show it in detail view
            if st.session_state.selected_doc:
                # Find the selected file object
                selected_file = next((f for f in uploaded_files if f.name == st.session_state.selected_doc), None)
                if selected_file:
                    st.markdown("---")
                    st.subheader(f"📄 {selected_file.name}")
                    
                    # Extract text based on file type
                    file_bytes = selected_file.getvalue()
                    file_extension = os.path.splitext(selected_file.name)[1].lower()
                    
                    if file_extension == ".docx":
                        text_content = extract_text_from_docx(file_bytes)
                    elif file_extension == ".pdf":
                        text_content = extract_text_from_pdf(file_bytes)
                    else:
                        text_content = "Unsupported file type"
                    
                    # Display stats
                    if file_extension == ".docx":
                        try:
                            doc = Document(BytesIO(file_bytes))
                            p_count, w_count, c_count = get_document_stats(doc)
                            
                            col1, col2, col3 = st.columns(3)
                            col1.metric("Paragraphs", p_count)
                            col2.metric("Words", w_count)
                            col3.metric("Characters", c_count)
                        except:
                            st.warning("Could not calculate document statistics")
                    elif file_extension == ".pdf":
                        p_count, w_count, c_count = get_pdf_stats(text_content)
                        
                        col1, col2, col3 = st.columns(3)
                        col1.metric("Lines", p_count)
                        col2.metric("Words", w_count)
                        col3.metric("Characters", c_count)
                    
                    # Display content
                    # Initialize session state for this document's content
                    content_key = f"doc_content_{selected_file.name}"
                    if content_key not in st.session_state:
                        st.session_state[content_key] = text_content
                    
                    # Update session state when text area content changes
                    edited_content = st.text_area(
                        "Document Text (Editable)",
                        value=st.session_state[content_key],
                        height=400,
                        key=f"edit_area_{selected_file.name}"
                    )
                    
                    # Update session state with current content (this ensures auto-save works)
                    st.session_state[content_key] = edited_content
                    
                    # Add explicit save button
                    if st.button("💾 Save Changes", key=f"save_{selected_file.name}"):
                        st.session_state[content_key] = edited_content
                        st.success("Changes saved successfully!")
                    
                    # Download buttons
                    st.markdown("---")
                    st.subheader("Save Options")
                    
                    if file_extension == ".docx":
                        # Create updated DOCX from edited content
                        updated_doc = create_document_from_text(edited_content)
                        updated_bytes = document_to_bytes(updated_doc)
                        
                        col1, col2, col3 = st.columns(3)
                        col1.download_button(
                            label="📄 Save as DOCX",
                            data=updated_bytes,
                            file_name=f"{os.path.splitext(selected_file.name)[0]}_edited.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            help="Downloads to your browser's default download folder"
                        )
                        
                        col2.download_button(
                            label="💾 Save as Text",
                            data=edited_content,
                            file_name=f"{os.path.splitext(selected_file.name)[0]}_edited.txt",
                            mime="text/plain",
                            help="Downloads to your browser's default download folder"
                        )
                        
                        # Create PDF from edited content
                        pdf_bytes = create_pdf_from_text(edited_content)
                        
                        col3.download_button(
                            label="🖨️ Save as PDF",
                            data=pdf_bytes,
                            file_name=f"{os.path.splitext(selected_file.name)[0]}_edited.pdf",
                            mime="application/pdf",
                            help="Downloads as a PDF document"
                        )
                    
                    else:  # PDF
                        # For PDFs, we can only save as text or convert to DOCX
                        new_doc = create_document_from_text(edited_content)
                        new_doc_bytes = document_to_bytes(new_doc)
                        
                        col1, col2 = st.columns(2)
                        col1.download_button(
                            label="📄 Save as DOCX",
                            data=new_doc_bytes,
                            file_name=f"{os.path.splitext(selected_file.name)[0]}_converted.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            help="Converts PDF content to Word document"
                        )
                        
                        col2.download_button(
                            label="💾 Save as Text",
                            data=edited_content,
                            file_name=f"{os.path.splitext(selected_file.name)[0]}_converted.txt",
                            mime="text/plain",
                            help="Extracted text from PDF"
                        )
                    
                    st.info("📁 Files save to your browser's default download location")
                    
                    # Button to go back to card view
                    if st.button("Back to Document Cards"):
                        st.session_state.selected_doc = None
                        st.rerun()
        else:
            st.info("Please upload .docx or .pdf files to get started. Documents will appear as cards below.")

    # Tab 2: Create New Document
    with tabs[1]:
        st.subheader("Create and Edit New Document")
        
        # Initialize session state for document content
        if "doc_content" not in st.session_state:
            st.session_state.doc_content = "# My Document\n\nStart typing here...\n\n- Item 1\n- Item 2\n\n**Bold text** and *italic text*"
        
        # Text area for document content
        st.session_state.doc_content = st.text_area(
            "Document Content (Markdown-style supported)",
            value=st.session_state.doc_content,
            height=400,
            key="doc_editor"
        )
        
        # Instructions
        with st.expander("💡 Formatting Tips"):
            st.markdown("""
            - **Headings**: Start lines with `#` (H1), `##` (H2), etc.
            - **Bold**: Wrap text with `**bold**`
            - **Italic**: Wrap text with `*italic*`
            - **Lists**: Use `- ` for bullet points
            - **New Paragraphs**: Separate with blank lines
            """)
        
        # Preview section
        st.subheader("Preview")
        st.markdown(st.session_state.doc_content)
        
        # Create document from content
        doc = create_document_from_text(st.session_state.doc_content)
        doc_bytes = document_to_bytes(doc)
        
        # Download buttons
        st.markdown("---")
        st.subheader("Save Your Document")
        
        col1, col2, col3 = st.columns(3)
        col1.download_button(
            label="📄 Save as DOCX",
            data=doc_bytes,
            file_name="my_document.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            help="Standard Word document format"
        )
        
        col2.download_button(
            label="💾 Save as Text",
            data=st.session_state.doc_content,
            file_name="my_document.txt",
            mime="text/plain",
            help="Plain text format"
        )
        
        # Create PDF from content
        pdf_bytes = create_pdf_from_text(st.session_state.doc_content)
        
        col3.download_button(
            label="🖨️ Save as PDF",
            data=pdf_bytes,
            file_name="my_document.pdf",
            mime="application/pdf",
            help="Downloads as a PDF document"
        )
        
        st.info("📁 Files save to your browser's default download location")

    # Additional information
    with st.expander("ℹ️ How to use this tool"):
        st.markdown("""
        1. **Supported Formats**:
           - Microsoft Word (.docx files)
           - PDF Documents (.pdf files)
        2. **Upload Documents**: 
           - Use the file uploader to select documents
           - View documents as cards in the grid layout
           - Click "View/Edit" on any card to open the document
           - Save modifications using the download buttons
        3. **Create New Documents**:
           - Switch to the "Create New Document" tab
           - Edit content directly in the text area
           - See real-time preview of your document
           - Save when finished using the save buttons
        4. **File Saving**:
           - Files download directly to your browser's default download location
           - Check your Downloads folder or browser download bar
           - Some browsers may ask before downloading multiple files
        """)
        
        st.markdown("""
        **Note on PDF Handling**: 
        - PDF text extraction may not preserve original formatting
        - Complex layouts, images, and tables may not be fully extracted
        - Converted PDFs become editable text documents
        """)

render_write_prisoners()
